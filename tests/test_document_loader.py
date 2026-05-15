from io import BytesIO

from docx import Document

from gauntlet_core.document_loader import extract_text_from_bytes, load_document_from_bytes


def test_extract_txt():
    assert extract_text_from_bytes("paper.txt", b"The paper resolves the anomaly.") == "The paper resolves the anomaly."


def test_load_txt_creates_source_spans():
    loaded = load_document_from_bytes(
        "paper.txt",
        b"The paper resolves the anomaly because the mechanism is measured. Evidence shows 12 tests.",
    )

    assert loaded.text.startswith("The paper resolves")
    assert loaded.source_spans
    assert loaded.source_spans[0].anchor_id == "S1"
    assert loaded.source_spans[0].page_number is None
    assert loaded.source_spans[0].char_start == 0


def test_extract_markdown():
    assert "resolves" in extract_text_from_bytes("paper.md", b"# Title\n\nThis resolves the issue.")


def test_load_markdown_creates_source_spans():
    loaded = load_document_from_bytes(
        "paper.md",
        b"# Abstract\n\nThe framework resolves the paradox through a mechanism. Evidence reports 42 tests.",
    )

    assert loaded.source_spans
    assert loaded.source_spans[0].char_start >= 0


def test_extract_docx():
    document = Document()
    document.add_paragraph("The framework explains the anomaly through a mechanism.")
    buffer = BytesIO()
    document.save(buffer)

    text = extract_text_from_bytes("paper.docx", buffer.getvalue())

    assert "explains the anomaly" in text


def test_load_docx_creates_source_spans():
    document = Document()
    document.add_paragraph("The framework explains the anomaly through a mechanism.")
    document.add_paragraph("Smith et al. (2024) reports 12 measurements.")
    buffer = BytesIO()
    document.save(buffer)

    loaded = load_document_from_bytes("paper.docx", buffer.getvalue())

    assert "explains the anomaly" in loaded.text
    assert loaded.source_spans
    assert loaded.source_spans[0].page_number is None


def test_extract_pdf_smoke():
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 300 144] /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length 58 >> stream\n"
        b"BT /F1 12 Tf 72 100 Td (The paper resolves the issue.) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n"
        b"0000000058 00000 n \n0000000115 00000 n \n0000000251 00000 n \n"
        b"0000000321 00000 n \ntrailer << /Root 1 0 R /Size 6 >>\nstartxref\n430\n%%EOF\n"
    )

    text = extract_text_from_bytes("paper.pdf", pdf_bytes)

    assert "resolves" in text


def test_load_pdf_preserves_page_numbers():
    pdf_bytes = (
        b"%PDF-1.4\n"
        b"1 0 obj << /Type /Catalog /Pages 2 0 R >> endobj\n"
        b"2 0 obj << /Type /Pages /Kids [3 0 R] /Count 1 >> endobj\n"
        b"3 0 obj << /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 300 144] /Contents 5 0 R >> endobj\n"
        b"4 0 obj << /Type /Font /Subtype /Type1 /BaseFont /Helvetica >> endobj\n"
        b"5 0 obj << /Length 58 >> stream\n"
        b"BT /F1 12 Tf 72 100 Td (The paper resolves the issue.) Tj ET\n"
        b"endstream endobj\n"
        b"xref\n0 6\n0000000000 65535 f \n0000000009 00000 n \n"
        b"0000000058 00000 n \n0000000115 00000 n \n0000000251 00000 n \n"
        b"0000000321 00000 n \ntrailer << /Root 1 0 R /Size 6 >>\nstartxref\n430\n%%EOF\n"
    )

    loaded = load_document_from_bytes("paper.pdf", pdf_bytes)

    assert loaded.source_spans
    assert loaded.source_spans[0].page_number == 1
    assert "resolves" in loaded.source_spans[0].text
