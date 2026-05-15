from __future__ import annotations

from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
import re

from .models import SourceSpan


SUPPORTED_EXTENSIONS = {".txt", ".md", ".pdf", ".docx"}


@dataclass(frozen=True)
class LoadedDocument:
    filename: str
    text: str
    source_spans: list[SourceSpan]


def extract_text_from_path(path: str | Path) -> str:
    path = Path(path)
    return extract_text_from_bytes(path.name, path.read_bytes())


def load_document_from_path(path: str | Path) -> LoadedDocument:
    path = Path(path)
    return load_document_from_bytes(path.name, path.read_bytes())


def extract_text_from_bytes(filename: str, data: bytes) -> str:
    return load_document_from_bytes(filename, data).text


def load_document_from_bytes(filename: str, data: bytes) -> LoadedDocument:
    extension = Path(filename).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))
        raise ValueError(f"Unsupported file type '{extension}'. Supported types: {supported}")
    if extension in {".txt", ".md"}:
        text = decode_text(data)
        return LoadedDocument(filename, text, build_source_spans(text))
    if extension == ".pdf":
        return extract_pdf_document(filename, data)
    if extension == ".docx":
        text = extract_docx_text(data)
        return LoadedDocument(filename, text, build_source_spans(text))
    raise ValueError(f"Unsupported file type '{extension}'.")


def decode_text(data: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return data.decode(encoding)
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace")


def extract_pdf_text(data: bytes) -> str:
    return extract_pdf_document("uploaded.pdf", data).text


def extract_pdf_document(filename: str, data: bytes) -> LoadedDocument:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise RuntimeError("PDF support requires the pypdf package. Run pip install -r requirements.txt.") from exc

    reader = PdfReader(BytesIO(data))
    text_parts: list[str] = []
    spans: list[SourceSpan] = []
    offset = 0
    for page_index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if not page_text.strip():
            continue
        if text_parts:
            text_parts.append("\n")
            offset += 1
        text_parts.append(page_text)
        page_spans = build_source_spans(page_text, page_number=page_index, offset=offset, start_index=len(spans) + 1)
        spans.extend(page_spans)
        offset += len(page_text)
    text = "".join(text_parts)
    return LoadedDocument(filename, text, spans)


def extract_docx_text(data: bytes) -> str:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("DOCX support requires the python-docx package. Run pip install -r requirements.txt.") from exc

    document = Document(BytesIO(data))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    table_cells: list[str] = []
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_cells.append(cell.text.strip())
    return "\n".join(paragraphs + table_cells)


def build_source_spans(
    text: str,
    page_number: int | None = None,
    offset: int = 0,
    start_index: int = 1,
) -> list[SourceSpan]:
    spans: list[SourceSpan] = []
    for match in re.finditer(r"[^.!?\n][^.!?]*(?:[.!?]|$)", text):
        sentence = re.sub(r"\s+", " ", match.group(0)).strip()
        if len(sentence) < 20:
            continue
        sentence_index = start_index + len(spans)
        spans.append(
            SourceSpan(
                anchor_id=f"S{sentence_index}",
                section="Document",
                sentence_index=sentence_index,
                page_number=page_number,
                char_start=offset + match.start(),
                char_end=offset + match.end(),
                text=sentence,
            )
        )
    return spans
